"""
Generador del documento Word profesional Safe Link — Manual de Instalación.

Uso:
    cd docs
    python generate_docx.py

Produce: docs/Safe_Link_Manual_Instalacion.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches

# ─────────────────────────────────────────────────────────────────
# Paleta corporativa Safe Link
# ─────────────────────────────────────────────────────────────────
COLOR_ACCENT      = RGBColor(0x25, 0x63, 0xEB)  # Azul primario
COLOR_TEAL        = RGBColor(0x14, 0xB8, 0xA6)  # Teal secundario
COLOR_TEXT        = RGBColor(0x1F, 0x29, 0x37)  # Casi negro
COLOR_TEXT_MUTED  = RGBColor(0x52, 0x52, 0x5B)  # Gris medio
COLOR_TEXT_FAINT  = RGBColor(0x9C, 0xA3, 0xAF)  # Gris claro
COLOR_BG_SOFT     = RGBColor(0xF4, 0xF4, 0xF5)  # Background sutil
COLOR_BORDER      = RGBColor(0xE5, 0xE7, 0xEB)
COLOR_GREEN       = RGBColor(0x22, 0xC5, 0x5E)
COLOR_RED         = RGBColor(0xEF, 0x44, 0x44)
COLOR_AMBER       = RGBColor(0xEA, 0xB3, 0x08)

FONT_BODY    = "Inter"           # cuerpo
FONT_HEADING = "Calibri Light"   # fallback ampliamente disponible
FONT_MONO    = "Consolas"        # mono fallback (JetBrains Mono no es default Windows)


# ─────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex="E5E7EB", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def add_heading(doc, text, level=1, color=None):
    """Heading custom (no usa styles built-in para control fino)."""
    p = doc.add_paragraph()
    if level == 0:  # Title
        run = p.add_run(text)
        run.font.name = FONT_HEADING
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_ACCENT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
    elif level == 1:
        run = p.add_run(text)
        run.font.name = FONT_HEADING
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_TEXT
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run = p.add_run(text)
        run.font.name = FONT_HEADING
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_ACCENT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run = p.add_run(text)
        run.font.name = FONT_HEADING
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_TEXT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, size=10.5, italic=False, color=None, bold=False, mono=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_MONO if mono else FONT_BODY
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEXT
    return p


def add_callout(doc, text, kind="info"):
    """Caja de info/warn/note coloreada con borde lateral."""
    colors = {
        "info":  ("DBEAFE", "2563EB"),  # azul soft / azul borde
        "warn":  ("FEF3C7", "EAB308"),  # ámbar soft / ámbar
        "ok":    ("DCFCE7", "22C55E"),
        "error": ("FEE2E2", "EF4444"),
    }
    bg, border = colors.get(kind, colors["info"])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)

    # Borde lateral grueso a la izquierda
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), border)
    borders.append(left)
    for side in ("top", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), bg)
        borders.append(b)
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # margen vertical


def add_table(doc, headers, rows, col_widths_cm=None):
    """Tabla estilizada estilo design system."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1F2937")  # dark
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = FONT_HEADING
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, "F9FAFB")  # zebra
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_val))
            run.font.name = FONT_BODY
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_TEXT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for r in table.rows:
                r.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_code(doc, code: str, lang_hint: str = ""):
    """Bloque de codigo monoespaciado con fondo gris."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F4F4F5")
    set_cell_border(cell, color_hex="E5E7EB")

    p = cell.paragraphs[0]
    if lang_hint:
        eyebrow = p.add_run(f"{lang_hint}\n")
        eyebrow.font.name = FONT_BODY
        eyebrow.font.size = Pt(8)
        eyebrow.font.color.rgb = COLOR_TEXT_FAINT
        eyebrow.font.bold = True

    run = p.add_run(code)
    run.font.name = FONT_MONO
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()


def add_pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ─────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Margenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    # Estilo default
    styles = doc.styles
    style = styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT

    # ─── PORTADA ───────────────────────────────────────
    # Banner superior
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover_table.cell(0, 0)
    set_cell_bg(cell, "070708")
    set_cell_border(cell, color_hex="070708", size=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SAFE LINK")
    run.font.name = FONT_HEADING
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(0)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("M O N I T O R I N G")
    run.font.name = FONT_HEADING
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEAL
    p2.paragraph_format.space_after = Pt(40)

    # Titulo del documento
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual de Instalación\ny Puesta en Marcha")
    run.font.name = FONT_HEADING
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plataforma SaaS B2B · Asistencia biométrica multi-tenant")
    run.font.name = FONT_BODY
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = COLOR_TEXT_MUTED
    p.paragraph_format.space_after = Pt(80)

    # Metadata
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("Versión del manual",  "1.0 · 2026-05-20"),
        ("Aplica a Estación",    "v5.7+"),
        ("Aplica a Panel web",   "v0.7+"),
        ("Audiencia",            "Administradores y operadores de Safe Link"),
    ]
    for i, (k, v) in enumerate(meta):
        c1 = meta_table.rows[i].cells[0]
        c2 = meta_table.rows[i].cells[1]
        set_cell_border(c1, color_hex="FFFFFF", size=0)
        set_cell_border(c2, color_hex="FFFFFF", size=0)
        c1.width = Cm(5)
        c2.width = Cm(10)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(k)
        run.font.name = FONT_BODY
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_TEXT_FAINT
        p2 = c2.paragraphs[0]
        run = p2.add_run(v)
        run.font.name = FONT_BODY
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_TEXT

    add_pagebreak(doc)

    # ─── INDICE ────────────────────────────────────────
    add_heading(doc, "Contenido", level=1)
    toc_items = [
        ("1.",  "Qué es Safe Link"),
        ("2.",  "Antes de empezar — requisitos"),
        ("3.",  "Instalación del Panel Web"),
        ("4.",  "Instalación de la Estación física"),
        ("5.",  "Verificación end-to-end"),
        ("6.",  "Operación diaria"),
        ("7.",  "Actualizaciones"),
        ("8.",  "Desinstalación"),
        ("9.",  "Solución de problemas"),
        ("10.", "Soporte"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  ")
        run.font.name = FONT_HEADING
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_ACCENT
        run2 = p.add_run(title)
        run2.font.name = FONT_BODY
        run2.font.size = Pt(11)
        run2.font.color.rgb = COLOR_TEXT
        p.paragraph_format.space_after = Pt(3)
    add_pagebreak(doc)

    # ─── 1. QUE ES SAFE LINK ────────────────────────────
    add_heading(doc, "1. ¿Qué es Safe Link?", level=1)
    add_body(doc,
        "Safe Link Monitoring es una plataforma SaaS B2B para control biométrico de "
        "asistencia laboral. Cada empresa cliente instala estaciones físicas en sus "
        "sucursales (una PC con cámara web) que reconocen a los empleados por su "
        "rostro y registran sus marcaciones de entrada/salida en la nube. El "
        "administrador gestiona empleados, sucursales y reportes desde un panel web "
        "sin necesidad de tocar bases de datos.")

    add_heading(doc, "Arquitectura simplificada", level=2)
    add_code(doc,
        "┌──────────────────┐       ┌──────────────────┐       ┌──────────────────┐\n"
        "│ Estación física  │──────▶│     Supabase     │◀──────│   Panel web      │\n"
        "│ (PC + cámara)    │ REST  │  (cloud DB)      │  SSR  │ panel.safelink   │\n"
        "│ Python + React   │       │  PostgreSQL +    │       │  Multi-tenant    │\n"
        "│                  │       │  RLS + Realtime  │       │  Next.js 15      │\n"
        "└──────────────────┘       └──────────────────┘       └──────────────────┘\n"
        "        ▲                                                       │\n"
        "        │ heartbeat cada 60s                                    │\n"
        "        └─── estado online/offline visible en tiempo real ──────┘",
        lang_hint="ARQUITECTURA")

    add_heading(doc, "Roles del sistema", level=2)
    add_table(doc,
        headers=["Rol", "Quién", "Qué hace"],
        rows=[
            ["Empleado",      "Colaborador de la empresa",   "Se para frente a la estación → marca entrada/salida"],
            ["Administrador", "Dueño/RH de la empresa",      "Crea empleados, sucursales y estaciones desde el panel"],
            ["Superadmin",    "Equipo Safe Link",            "Soporte, monitoreo cross-empresa, releases"],
        ],
        col_widths_cm=[3, 4, 9])

    add_pagebreak(doc)

    # ─── 2. REQUISITOS ────────────────────────────────────
    add_heading(doc, "2. Antes de empezar — requisitos", level=1)

    add_heading(doc, "Para el Panel web", level=2)
    add_table(doc,
        headers=["Requisito", "Detalle"],
        rows=[
            ["Navegador",   "Chrome 120+, Edge 120+, Firefox 121+, Safari 17+"],
            ["Internet",    "Estable (≥5 Mbps)"],
            ["Cuenta",      "Email + contraseña entregados por Safe Link (o link de onboarding)"],
            ["Dispositivo", "Cualquier laptop/desktop. Resolución mínima recomendada: 1366×768"],
        ],
        col_widths_cm=[4, 12])

    add_heading(doc, "Para la Estación física", level=2)
    add_table(doc,
        headers=["Requisito", "Detalle"],
        rows=[
            ["Sistema operativo", "Windows 10 / 11 (64 bits)"],
            ["Cámara",            "Webcam USB o integrada (resolución mínima 720p)"],
            ["Procesador",        "Intel i3 8va gen o equivalente AMD (para inferencia ML local)"],
            ["RAM",               "4 GB mínimo, 8 GB recomendado"],
            ["Espacio",           "280 MB libres en disco"],
            ["Internet",          "Estable para sincronizar (la estación tiene buffer offline)"],
            ["Permisos",          "Administrador local en la PC para instalar"],
            ["Instalador",        "SafeLinkStation_Setup.exe (descarga de GitHub Releases)"],
        ],
        col_widths_cm=[4.5, 11.5])

    add_pagebreak(doc)

    # ─── 3. PANEL WEB ─────────────────────────────────────
    add_heading(doc, "3. Instalación del Panel Web", level=1)
    add_callout(doc,
        "No requiere instalación de software — es una aplicación web. "
        "Solo necesitas un navegador moderno.", kind="info")

    add_heading(doc, "Paso 3.1 — Recibir credenciales", level=2)
    add_body(doc,
        "El equipo de Safe Link te proporciona uno de estos dos accesos:")

    add_heading(doc, "Opción A: Credenciales directas (empresa ya creada)", level=3)
    add_code(doc,
        "URL:        https://panel.safelink.app\n"
        "Email:      admin@tuempresa.com\n"
        "Password:   (entregada en sobre sellado o por canal seguro)")

    add_heading(doc, "Opción B: Link de onboarding (empresa nueva — zero-touch)", level=3)
    add_code(doc, "https://panel.safelink.app/activar?token=ABC-123-XYZ")

    add_heading(doc, "Paso 3.2 — Primer ingreso", level=2)
    add_bullet(doc, "Abre el navegador en la URL proporcionada")
    add_bullet(doc, "Si tienes credenciales: ingresa email + password en el form de login")
    add_bullet(doc, "Si tienes link de onboarding: llena el wizard (nombre empresa, email admin, "
                    "password, primera sucursal)")
    add_bullet(doc, "Al hacer login serás redirigido a /tablero (vista principal)")

    add_heading(doc, "Paso 3.3 — Mapa de secciones del panel", level=2)
    add_table(doc,
        headers=["Grupo", "Sección", "Para qué"],
        rows=[
            ["Operación", "Inicio (/tablero)",        "Wall en vivo de tus estaciones"],
            ["Operación", "Empleados",                "Crear y gestionar personas"],
            ["Operación", "Sucursales",               "Locales físicos con ubicación en mapa"],
            ["Operación", "Estaciones",               "Las PCs físicas con cámara"],
            ["Análisis",  "Reportes",                 "Horas trabajadas, retardos"],
            ["Análisis",  "Ejecutivo",                "KPIs para gerencia"],
            ["Análisis",  "Actividad",                "Audit log"],
            ["Sistema",   "Notificaciones",           "Histórico de eventos"],
            ["Sistema",   "Configuración",            "Datos de la empresa"],
        ],
        col_widths_cm=[2.5, 4, 9.5])

    add_callout(doc,
        "Búsqueda global: presiona Cmd+K (Mac) o Ctrl+K (Windows) para abrir la "
        "paleta de comandos y navegar rápido.", kind="info")

    add_heading(doc, "Paso 3.4 — Configurar primera sucursal con ubicación", level=2)
    add_bullet(doc, "Sidebar → Sucursales → botón + Nueva sucursal")
    add_bullet(doc, "Tab Información: nombre, dirección")
    add_bullet(doc, "Tab Horario: hora apertura, cierre, tolerancia (minutos)")
    add_bullet(doc, "Tab Ubicación: click en \"Mi ubicación\" o click en el mapa para colocar el pin")
    add_bullet(doc, "Arrastra el pin para afinar la posición")
    add_bullet(doc, "Click Crear sucursal")
    add_bullet(doc, "Verifica en sidebar → Mapa que el pin aparezca con halo cyan pulsando")

    add_heading(doc, "Paso 3.5 — Crear empleados", level=2)
    add_bullet(doc, "Sidebar → Empleados → + Nuevo empleado")
    add_bullet(doc, "Nombre, apellido, código de empleado, sucursal asignada")
    add_bullet(doc, "Sube una foto frontal clara (buena iluminación, sin lentes oscuros)")
    add_bullet(doc, "Click Guardar")

    add_callout(doc,
        "La foto se sube a Supabase Storage y la estación la descargará automáticamente "
        "en su próxima sincronización para generar los embeddings faciales.", kind="ok")

    add_pagebreak(doc)

    # ─── 4. ESTACION FISICA ───────────────────────────────
    add_heading(doc, "4. Instalación de la Estación física", level=1)

    add_heading(doc, "Paso 4.1 — Descargar el instalador", level=2)
    add_body(doc,
        "Desde el equipo donde vas a instalar la estación, ve a:")
    add_code(doc, "https://github.com/Safe-LM/app-login-trabajadores-desktop/releases/latest")
    add_body(doc, "Descarga el archivo SafeLinkStation_Setup.exe (~280 MB). Guárdalo en Descargas.")

    add_heading(doc, "Paso 4.2 — Generar la API Key en el panel", level=2)
    add_body(doc,
        "Antes de ejecutar el instalador, necesitas crear la estación en el panel "
        "web para obtener su API Key:")
    add_bullet(doc, "Panel web → sidebar → Estaciones → botón + Registrar estación")
    add_bullet(doc, "Llena: nombre (ej: Recepción Norte) + sucursal asignada")
    add_bullet(doc, "Click Registrar")
    add_bullet(doc, "Copia la API Key que aparece (formato sk_xxxxxxxxxxxxxxxx)")

    add_callout(doc,
        "Solo se muestra una vez. Guárdala en un lugar seguro temporalmente. "
        "Si la pierdes, puedes regenerarla desde el panel.", kind="warn")

    add_heading(doc, "Paso 4.3 — Ejecutar el instalador", level=2)
    add_bullet(doc, "Doble click en SafeLinkStation_Setup.exe")
    add_bullet(doc, "Si Windows muestra \"Editor desconocido\" (normal sin firma): "
                    "Más información → Ejecutar de todas formas")
    add_bullet(doc, "Acepta el aviso UAC → Sí")
    add_bullet(doc, "Click Siguiente en el wizard")

    add_heading(doc, "Paso 4.4 — Carpeta de instalación", level=2)
    add_body(doc,
        "Deja la ruta por defecto (C:\\Program Files\\Safe Link Station\\) y "
        "click Siguiente. Si tu disco C: tiene poco espacio, puedes cambiar a "
        "D:\\Safe Link Station\\.")

    add_heading(doc, "Paso 4.5 — Configuración inicial", level=2)
    add_body(doc,
        "Esta es la pantalla más importante. Llena con cuidado:")

    add_table(doc,
        headers=["Campo", "Valor"],
        rows=[
            ["Nombre de la estación",
             "El que tú quieras. Descriptivo del lugar (ej: Recepción, Almacén Norte)."],
            ["API Key",
             "La que copiaste en el paso 4.2 (empieza con sk_)"],
            ["URL Supabase",
             "Viene precargada. Si te dieron credenciales personalizadas, pégalas."],
            ["Anon Key",
             "Viene precargada igual que URL."],
            ["Iniciar con Windows",
             "Marca esta casilla si la PC se usa 24/7 (recomendado para producción)"],
        ],
        col_widths_cm=[5, 11])

    add_body(doc, "Click Siguiente → Finalizar.")

    add_heading(doc, "Paso 4.6 — Primera ejecución de la estación", level=2)
    add_body(doc, "Al finalizar, la estación se abre automáticamente en pantalla completa:")
    add_bullet(doc, "Splash con logo Safe Link mientras carga (~5 segundos)")
    add_bullet(doc, "Pide acceso a la cámara → acepta cuando Windows pregunte")
    add_bullet(doc, "Sincroniza la lista de empleados con Supabase (~30 segundos)")
    add_bullet(doc, "Por cada empleado con foto, genera 10 embeddings faciales y los sube a la nube")
    add_bullet(doc, "Cuando aparece \"Listo — buscando rostro…\", la estación está activa")

    add_callout(doc,
        "La primera sincronización puede tardar 1-2 minutos si hay muchos empleados. "
        "Verás \"Sincronizando...\" en la parte superior derecha.", kind="info")

    add_pagebreak(doc)

    # ─── 5. VERIFICACION ───────────────────────────────────
    add_heading(doc, "5. Verificación end-to-end", level=1)

    add_heading(doc, "Checklist en la estación", level=2)
    add_bullet(doc, "☐  Logo SAFE LINK MONITORING visible arriba a la izquierda")
    add_bullet(doc, "☐  Cámara muestra video en vivo (no pantalla negra)")
    add_bullet(doc, "☐  Indicador \"EN LÍNEA\" verde en esquina superior derecha")
    add_bullet(doc, "☐  Reloj en vivo actualizándose cada segundo")
    add_bullet(doc, "☐  Mensaje \"Buscando rostro...\" abajo de la cámara")
    add_bullet(doc, "☐  Al pararse frente a la cámara: detecta el rostro (cuadro rojo/verde)")

    add_heading(doc, "Checklist en el panel web", level=2)
    add_bullet(doc, "Sidebar → Tablero → debes ver la nueva estación como tile con border verde")
    add_bullet(doc, "Sidebar → Estaciones → la estación aparece con badge \"● En línea\"")
    add_bullet(doc, "El heartbeat debe ser reciente (hace 12s, hace 30s, etc.)")
    add_bullet(doc, "Sidebar → Mapa → si la sucursal tiene ubicación, el pin debe ser verde")

    add_heading(doc, "Prueba de marcación", level=2)
    add_bullet(doc, "Para frente a la estación a un empleado con foto enrollada")
    add_bullet(doc, "La estación debe reconocerlo en <1 segundo y mostrar ✓ Bienvenido, [Nombre]")
    add_bullet(doc, "Suena la campanita 🔔")
    add_bullet(doc, "Panel web → Dashboard → el empleado aparece en la tabla de asistencias")
    add_bullet(doc, "Sidebar → Tablero → aside \"Marcaciones recientes\" actualiza en tiempo real")

    add_callout(doc,
        "Si todo lo anterior pasa, la instalación es exitosa.", kind="ok")

    add_pagebreak(doc)

    # ─── 6. OPERACION DIARIA ───────────────────────────────
    add_heading(doc, "6. Operación diaria", level=1)

    add_heading(doc, "Para el administrador (panel web)", level=2)
    add_bullet(doc, "Tablero: revisar al inicio del día que todas las estaciones estén En línea")
    add_bullet(doc, "Dashboard: ver KPIs del día (presentes, ausentes, asistencia %)")
    add_bullet(doc, "Notificaciones: revisar el bell 🔔 — alertas de estaciones offline o llegadas tarde")
    add_bullet(doc, "Empleados: agregar/desactivar empleados cuando cambien en la empresa")
    add_bullet(doc, "Reportes: exportar a Excel al final del mes")

    add_heading(doc, "Para los empleados (estación)", level=2)
    add_bullet(doc, "Solo pararse frente a la cámara con la cara visible")
    add_bullet(doc, "Esperar la confirmación visual + sonora (~1 segundo)")
    add_bullet(doc, "Si no reconoce: alejarse 1 metro y volver a acercarse")
    add_bullet(doc, "Si persiste: avisar al admin para verificar que su foto esté actualizada")

    add_heading(doc, "Eventos comunes y respuesta", level=2)
    add_table(doc,
        headers=["Evento", "Severidad", "Acción del admin"],
        rows=[
            ["Estación offline >5 min",      "Warn",     "Verificar conectividad y reinicio"],
            ["Estación offline >30 min",     "Crítico",  "Visita física a la sucursal"],
            ["Cámara con error",             "Error",    "Verificar USB, reiniciar estación"],
            ["Health score <50",             "Warn",     "Revisar logs en estación"],
            ["Llegada tarde de empleado",    "Info/Warn","Auto-notificado en bell del panel"],
        ],
        col_widths_cm=[5, 3, 8])

    add_pagebreak(doc)

    # ─── 7. ACTUALIZACIONES ────────────────────────────────
    add_heading(doc, "7. Actualizaciones", level=1)

    add_heading(doc, "Panel web", level=2)
    add_body(doc,
        "Se actualiza automáticamente — solo recarga la página (Ctrl+Shift+R). "
        "No requiere acción del cliente. Los releases se anuncian en el panel "
        "mediante banner superior.")

    add_heading(doc, "Estación", level=2)
    add_body(doc,
        "La estación detecta nuevas versiones al arrancar y muestra una notificación:")
    add_bullet(doc, "Aceptar → descarga e instala automáticamente (~2 min)")
    add_bullet(doc, "Rechazar → sigue con la versión actual; te avisará la próxima vez")

    add_body(doc, "Para desactivar auto-updates, edita el archivo:")
    add_code(doc, "C:\\Program Files\\Safe Link Station\\.env")
    add_body(doc, "Y agrega la línea:")
    add_code(doc, "AUTO_UPDATE_ENABLED=false")

    # ─── 8. DESINSTALACION ─────────────────────────────────
    add_heading(doc, "8. Desinstalación", level=1)

    add_heading(doc, "Estación", level=2)
    add_heading(doc, "Opción A — Desde menú Inicio", level=3)
    add_bullet(doc, "Menú Inicio → buscar Safe Link Station")
    add_bullet(doc, "Click derecho → Desinstalar")

    add_heading(doc, "Opción B — Desde Configuración de Windows", level=3)
    add_bullet(doc, "Configuración → Aplicaciones → Aplicaciones instaladas")
    add_bullet(doc, "Buscar Safe Link Station → ⋯ → Desinstalar")

    add_callout(doc,
        "Tus datos (asistencias offline en data/db/, logs) se conservan por seguridad. "
        "Si quieres borrarlo todo, elimina manualmente C:\\Program Files\\Safe Link "
        "Station\\ después de desinstalar.", kind="warn")

    add_heading(doc, "Panel web", level=2)
    add_body(doc,
        "No hay que desinstalar nada. Si una empresa quiere darse de baja, contacta "
        "al soporte de Safe Link para eliminar su tenant y datos.")

    add_pagebreak(doc)

    # ─── 9. SOLUCION DE PROBLEMAS ──────────────────────────
    add_heading(doc, "9. Solución de problemas", level=1)

    add_heading(doc, "Panel web", level=2)
    add_table(doc,
        headers=["Problema", "Solución"],
        rows=[
            ["Login OK pero /dashboard redirige a /login",
             "El middleware no encuentra empresa_id en el JWT — pide al admin que vuelva a crear tu usuario vía onboarding"],
            ["Tabla de empleados vacía aunque existen",
             "RLS está bloqueando — contacta soporte"],
            ["Mapa no carga (pantalla negra)",
             "Verifica internet — los tiles vienen de CartoDB CDN"],
            ["Pin de sucursal no aparece en /mapa",
             "La sucursal no tiene lat/lng — edita en /sucursales → tab Ubicación"],
            ["Notificaciones no llegan en tiempo real",
             "Verifica que Realtime esté habilitado en tu plan Supabase"],
        ],
        col_widths_cm=[6, 10])

    add_heading(doc, "Estación", level=2)
    add_table(doc,
        headers=["Problema", "Solución"],
        rows=[
            ["\"Editor desconocido\" al ejecutar el instalador",
             "Normal sin firma digital. Más información → Ejecutar de todas formas"],
            ["\"Cámara no disponible\"",
             "Cierra Teams/Zoom/Skype/cualquier app que use la cámara. Reinicia."],
            ["\"Estación sin conexión\" en el panel",
             "Verifica internet. La estación reintentará cada 60s."],
            ["El instalador se cierra sin avisar",
             "Ejecuta como administrador (click derecho → Ejecutar como administrador)"],
            ["No genera embeddings",
             "Verifica que los empleados tengan foto subida en el panel"],
            ["Cámara invertida o lateral",
             "Es una limitación de la cámara, no del software"],
            ["No reconoce empleados nuevos",
             "Espera 60s tras crearlos en el panel — sincronización automática"],
            ["Heartbeat se queda en \"hace 5 min\" o más",
             "La estación perdió internet o se cerró. Verifica que esté corriendo."],
        ],
        col_widths_cm=[6, 10])

    add_heading(doc, "Cómo editar la configuración después de instalar", level=2)
    add_body(doc, "Edita el archivo:")
    add_code(doc, "C:\\Program Files\\Safe Link Station\\.env")
    add_body(doc,
        "Abre con Bloc de notas como administrador (click derecho → \"Ejecutar como "
        "administrador\"). Modifica los valores y reinicia la estación.")

    add_callout(doc,
        "Si abres con doble click sin \"como administrador\", Windows no te dejará guardar.",
        kind="warn")

    add_pagebreak(doc)

    # ─── 10. SOPORTE ───────────────────────────────────────
    add_heading(doc, "10. Soporte", level=1)

    add_table(doc,
        headers=["Canal", "Para qué"],
        rows=[
            ["Email",         "soporte@safelink.app — soporte general, dudas, incidencias"],
            ["GitHub",        "github.com/Safe-LM/app-login-trabajadores-desktop — issues técnicos, request de features"],
            ["Logs estación", "C:\\Program Files\\Safe Link Station\\logs\\ — útil al reportar bugs"],
            ["Docs técnicas", "docs/arquitectura.md, docs/runbook-estacion.md, docs/runbook-panel.md"],
        ],
        col_widths_cm=[3, 13])

    add_heading(doc, "Datos a incluir en un ticket de soporte", level=2)
    add_bullet(doc, "ID de empresa (visible en el panel → Configuración)")
    add_bullet(doc, "Nombre de la estación afectada (si aplica)")
    add_bullet(doc, "Captura de pantalla del problema")
    add_bullet(doc, "Hora exacta del incidente")
    add_bullet(doc, "Logs (si es un problema de estación, adjuntar archivos de logs/)")

    # ─── FOOTER FINAL ──────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Fin del documento —")
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = COLOR_TEXT_FAINT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("© 2026 Safe Link Monitoring · Manual de Instalación v1.0")
    run.font.name = FONT_BODY
    run.font.size = Pt(8.5)
    run.font.color.rgb = COLOR_TEXT_FAINT

    # ─── GUARDAR ──────────────────────────────────────────
    out = Path(__file__).parent / "Safe_Link_Manual_Instalacion.docx"
    doc.save(out)
    print(f"[OK] Generado: {out}")
    print(f"     Tamano: {out.stat().st_size // 1024} KB")


if __name__ == "__main__":
    build()
